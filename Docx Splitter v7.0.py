# Docx Splitter  - 完整修复版
import os
import copy
import re
from docx import Document

def split_by_pages(file_path, pages_per_file=1):
    """按页数分割"""
    try:
        import win32com.client
    except ImportError:
        print("❌ 需要安装: pip install pywin32")
        return False
    
    base = os.path.splitext(file_path)[0]
    abs_path = os.path.abspath(file_path)
    output_dir = base + "_分割结果"
    
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    word = None
    doc = None
    
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        
        doc = word.Documents.Open(abs_path)
        doc.Repaginate()
        
        total_pages = doc.ComputeStatistics(2)
        
        print(f"\n📄 总页数: {total_pages}")
        print(f"📦 每 {pages_per_file} 页一个文件")
        
        expected_files = (total_pages + pages_per_file - 1) // pages_per_file
        print(f"📁 预计生成: {expected_files} 个文件")
        print(f"📂 输出目录: {output_dir}\n")
        
        file_counter = 0
        current_page = 1
        
        while current_page <= total_pages:
            end_page = min(current_page + pages_per_file - 1, total_pages)
            
            start_range = doc.GoTo(What=1, Which=1, Count=current_page)
            start_pos = start_range.Start
            
            if end_page < total_pages:
                end_range = doc.GoTo(What=1, Which=1, Count=end_page + 1)
                end_pos = end_range.Start
            else:
                end_pos = doc.Content.End
            
            copy_range = doc.Range(Start=start_pos, End=end_pos)
            copy_range.Copy()
            
            new_doc = word.Documents.Add()
            new_doc.Content.Paste()
            
            file_counter += 1
            output_name = f"第{current_page:03d}-{end_page:03d}页.docx"
            output_path = os.path.join(output_dir, output_name)
            
            new_doc.SaveAs2(os.path.abspath(output_path), 16)
            new_doc.Close(False)
            
            if file_counter % 20 == 0 or file_counter == expected_files:
                print(f"✔ 进度: {file_counter}/{expected_files}")
            
            current_page = end_page + 1
        
        print(f"\n✅ 完成！共生成 {file_counter} 个文件")
        print(f"📂 位置: {output_dir}")
        return True
        
    except Exception as e:
        print(f"\n❌ 错误: {e}")
        return False
        
    finally:
        try:
            if doc: doc.Close(False)
            if word: word.Quit()
        except:
            pass


def split_by_paragraphs(file_path, paras_per_file=1, skip_empty=True):
    """
    按段落分割 - 修复版
    skip_empty: 是否跳过空段落
    """
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"❌ 无法打开文件: {e}")
        return False
    
    base = os.path.splitext(file_path)[0]
    file_name = os.path.basename(base)
    output_dir = base + "_分割结果"
    
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # 收集所有段落元素
    all_para_elements = []
    
    for para in doc.paragraphs:
        # 如果跳过空段落
        if skip_empty and not para.text.strip():
            continue
        all_para_elements.append(para._element)
    
    total_paras = len(all_para_elements)
    
    if total_paras == 0:
        print("❌ 文档没有有效段落！")
        return False
    
    print(f"\n📝 有效段落数: {total_paras}")
    print(f"📦 每 {paras_per_file} 段一个文件")
    
    expected_files = (total_paras + paras_per_file - 1) // paras_per_file
    print(f"📁 预计生成: {expected_files} 个文件")
    print(f"📂 输出目录: {output_dir}\n")
    
    file_counter = 0
    
    # 分块处理
    for start in range(0, total_paras, paras_per_file):
        end = min(start + paras_per_file, total_paras)
        
        # 获取本次要处理的段落元素
        chunk_elements = all_para_elements[start:end]
        
        # 创建新文档
        new_doc = Document()
        
        # 删除新文档的默认空段落
        while len(new_doc.element.body):
            new_doc.element.body.remove(new_doc.element.body[0])
        
        # 深拷贝段落到新文档
        for elem in chunk_elements:
            new_elem = copy.deepcopy(elem)
            new_doc.element.body.append(new_elem)
        
        # 保存
        file_counter += 1
        output_name = f"{file_name}_{file_counter:04d}.docx"
        output_path = os.path.join(output_dir, output_name)
        
        try:
            new_doc.save(output_path)
        except Exception as e:
            print(f"❌ 保存失败 {output_name}: {e}")
            continue
        
        # 显示进度
        if file_counter % 50 == 0 or file_counter == expected_files:
            percent = int(file_counter / expected_files * 100)
            print(f"✔ 进度: {file_counter}/{expected_files} ({percent}%)")
    
    print(f"\n✅ 完成！共生成 {file_counter} 个文件")
    print(f"📂 位置: {output_dir}")
    return True


def detect_headings(doc):
    """检测标题"""
    headings = []
    
    for i, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else ""
        text = para.text.strip()
        
        if not text:
            continue
        
        heading_level = None
        
        if "Heading" in style_name or "标题" in style_name:
            match = re.search(r'(\d+)', style_name)
            heading_level = int(match.group(1)) if match else 1
        
        if heading_level is None:
            patterns = [
                (r'^第[一二三四五六七八九十百千\d]+[章节篇部]', 1),
                (r'^[一二三四五六七八九十]+[、.．]', 2),
                (r'^\d{1,2}[、.．]\s*\S', 2),
                (r'^[（(]\d+[)）]', 3),
            ]
            for pattern, level in patterns:
                if re.match(pattern, text):
                    heading_level = level
                    break
        
        if heading_level:
            headings.append({
                'index': i,
                'level': heading_level,
                'text': text[:40]
            })
    
    return headings


def smart_split(file_path, split_level=1):
    """智能分割"""
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"❌ 无法打开文件: {e}")
        return False
    
    base = os.path.splitext(file_path)[0]
    output_dir = base + "_分割结果"
    
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    headings = detect_headings(doc)
    paragraphs = list(doc.paragraphs)
    
    split_points = [h['index'] for h in headings if h['level'] <= split_level]
    
    if not split_points:
        print("❌ 没有检测到符合条件的标题")
        return False
    
    if split_points[0] != 0:
        split_points.insert(0, 0)
    
    split_points.append(len(paragraphs))
    
    expected_files = len(split_points) - 1
    print(f"\n📑 将分割为 {expected_files} 个文件")
    print(f"📂 输出目录: {output_dir}\n")
    
    file_counter = 0
    
    for i in range(len(split_points) - 1):
        start_idx = split_points[i]
        end_idx = split_points[i + 1]
        
        title = paragraphs[start_idx].text.strip()[:30]
        safe_title = re.sub(r'[\\/:*?"<>|\r\n]', '_', title)
        
        new_doc = Document()
        while len(new_doc.element.body):
            new_doc.element.body.remove(new_doc.element.body[0])
        
        for j in range(start_idx, end_idx):
            new_doc.element.body.append(
                copy.deepcopy(paragraphs[j]._element)
            )
        
        file_counter += 1
        output_name = f"{file_counter:03d}_{safe_title}.docx"
        output_path = os.path.join(output_dir, output_name)
        
        new_doc.save(output_path)
        print(f"✔ [{file_counter}] {output_name}")
    
    print(f"\n✅ 完成！共生成 {file_counter} 个文件")
    print(f"📂 位置: {output_dir}")
    return True


def get_doc_info(file_path):
    """获取文档信息"""
    info = {
        'pages': None,
        'paragraphs': 0,
        'headings': 0,
        'valid_paragraphs': 0
    }
    
    try:
        doc = Document(file_path)
        info['paragraphs'] = len(doc.paragraphs)
        info['headings'] = len(detect_headings(doc))
        info['valid_paragraphs'] = sum(
            1 for p in doc.paragraphs if p.text.strip()
        )
    except:
        pass
    
    # 尝试获取页数
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        d = word.Documents.Open(os.path.abspath(file_path))
        d.Repaginate()
        info['pages'] = d.ComputeStatistics(2)
        d.Close(False)
        word.Quit()
    except:
        pass
    
    return info


def process_one_file():
    """处理单个文件"""
    print("\n" + "─" * 50)
    path = input("📁 请输入 DOCX 文件路径 (拖入即可): ").strip().strip('"\'')
    
    if not path:
        print("❌ 未输入路径")
        return True  # 继续循环
    
    if not os.path.exists(path):
        print("❌ 文件不存在！")
        return True
    
    if not path.lower().endswith('.docx'):
        print("❌ 请输入 .docx 文件！")
        return True
    
    # 分析文档
    print("\n⏳ 正在分析文档...")
    info = get_doc_info(path)
    
    print("\n" + "─" * 50)
    print("📊 文档信息:")
    print("─" * 50)
    if info['pages']:
        print(f"   📄 页数: {info['pages']} 页")
    else:
        print(f"   📄 页数: 无法检测 (需要Word)")
    print(f"   📝 总段落: {info['paragraphs']} 段")
    print(f"   📝 有效段落: {info['valid_paragraphs']} 段 (非空)")
    print(f"   📑 检测标题: {info['headings']} 个")
    print("─" * 50)
    
    # 选择模式
    print("\n🔧 选择分割模式:")
    print("   [1] 📄 按页数分割 (需要Word)")
    print("   [2] 📝 按段落分割")
    print("   [3] 🧠 智能分割 (按标题)")
    print("   [0] ↩️  返回")
    
    mode = input("\n请选择 [0/1/2/3]: ").strip()
    
    if mode == "0":
        return True
    
    elif mode == "1":
        if not info['pages']:
            print("\n⚠️ 按页分割需要 Windows + Word + pywin32")
            confirm = input("是否继续? [y/N]: ").strip().lower()
            if confirm != 'y':
                return True
        
        n = input(f"\n每多少页一个文件? [默认1]: ").strip()
        n = int(n) if n.isdigit() and int(n) > 0 else 1
        split_by_pages(path, n)
    
    elif mode == "2":
        print(f"\n💡 有效段落数: {info['valid_paragraphs']}")
        n = input(f"每多少段一个文件? [默认1]: ").strip()
        n = int(n) if n.isdigit() and int(n) > 0 else 1
        
        skip = input("是否跳过空段落? [Y/n]: ").strip().lower()
        skip_empty = skip != 'n'
        
        split_by_paragraphs(path, n, skip_empty)
    
    elif mode == "3":
        if info['headings'] == 0:
            print("\n❌ 未检测到标题，无法智能分割")
            return True
        
        n = input(f"\n按几级标题分割? [默认1]: ").strip()
        n = int(n) if n.isdigit() and int(n) > 0 else 1
        smart_split(path, n)
    
    else:
        print("❌ 无效选择")
    
    return True


def main():
    """主程序 - 循环处理"""
    print("=" * 60)
    print("      📄 DOCX 无损分割工具 v6.2")
    print("=" * 60)
    
    while True:
        process_one_file()
        
        print("\n" + "=" * 50)
        print("继续操作:")
        print("   [1] 📂 分割下一个文件")
        print("   [0] 🚪 退出程序")
        print("=" * 50)
        
        choice = input("\n请选择 [1/0]: ").strip()
        
        if choice == "0" or choice.lower() == 'q':
            print("\n👋 再见！")
            break
        
        # 任何其他输入都继续


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n\n👋 程序已终止")
    except Exception as e:
        print(f"\n❌ 程序错误: {e}")
        input("\n按回车键退出...")